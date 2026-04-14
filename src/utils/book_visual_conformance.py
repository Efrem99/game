"""Book-driven visual conformance helpers for gameplay video runs."""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Set

from utils.video_bot_plan import build_video_bot_events, resolve_video_bot_plan_name


def _normalize_token(value: Any) -> str:
    token = str(value or "").strip().lower().replace("-", "_").replace(" ", "_")
    token = token.replace("__", "_")
    return token


def _normalize_text(value: Any) -> str:
    text = str(value or "").lower().replace("ё", "е")
    return " ".join(text.split())


_RUNTIME_TO_CANONICAL = {
    "town": "town_center",
    "town_center": "town_center",
    "castle": "castle_interior",
    "castle_interior": "castle_interior",
    "prince_chamber": "prince_chamber",
    "world_map_gallery": "world_map_gallery",
    "royal_laundry": "royal_laundry",
    "throne_hall": "throne_hall",
    "port": "port_market",
    "docks": "port_market",
    "boats": "port_market",
    "port_town": "port_market",
    "port_market_walk": "port_market",
    "port_market_memory": "port_market",
    "city_waterfront": "port_market",
    "kremor_forest": "krimora_forest",
    "kremor_forest_crash": "krimora_forest",
    "kremor_forest_cage": "krimora_forest",
    "krimora_forest": "krimora_forest",
    "krimora_forest_crash": "krimora_forest",
    "krimora_forest_cage": "krimora_forest",
    "dwarven_caves": "dwarven_caves",
    "dwarven_caves_gate": "dwarven_caves",
    "dwarven_caves_halls": "dwarven_caves",
    "dwarven_caves_throne": "dwarven_caves",
    "sharuan_forest_bridge": "sharuan_forest_bridge",
    "paradise_vision": "paradise_vision",
    "old_forest": "old_forest",
    "training_grounds": "training_grounds",
}


BOOK_LOCATION_PROFILES: Dict[str, Dict[str, Any]] = {
    "town_center": {
        "title": "Town Center",
        "aliases": ["town center", "sharuan square", "town hall", "городская площадь", "центр города"],
        "cues": ["square", "town hall", "fountain", "merchant", "площадь", "ратуш", "фонтан", "рынок"],
        "min_cues": 1,
    },
    "port_market": {
        "title": "Sharuan Port Market",
        "aliases": [
            "sharuan port market",
            "port market",
            "adalin market",
            "порт",
            "рынок",
            "шаруан",
        ],
        "cues": [
            "waterfront",
            "dock",
            "boat",
            "ship",
            "пристан",
            "причал",
            "корабл",
            "лавк",
            "торгов",
        ],
        "min_cues": 2,
    },
    "castle_interior": {
        "title": "Castle Interior",
        "aliases": ["castle interior", "castle", "замок", "дворец"],
        "cues": ["corridor", "hall", "interior", "коридор", "внутри", "двор"],
        "min_cues": 1,
    },
    "prince_chamber": {
        "title": "Prince Chamber",
        "aliases": ["prince chamber", "chamber", "покои", "комната принца"],
        "cues": ["map", "private room", "косм", "покои", "комната"],
        "min_cues": 1,
    },
    "world_map_gallery": {
        "title": "World Map Gallery",
        "aliases": ["world map", "map gallery", "карта мира", "галерея карт"],
        "cues": ["lorena", "argelia", "doren", "zarandia", "krimora", "лоре", "дорен", "заранд"],
        "min_cues": 1,
    },
    "royal_laundry": {
        "title": "Royal Laundry",
        "aliases": ["royal laundry", "laundry", "прачеч", "гардероб"],
        "cues": ["outfit", "wardrobe", "juice", "одежд", "переод", "пятн"],
        "min_cues": 1,
    },
    "throne_hall": {
        "title": "Throne Hall",
        "aliases": ["throne hall", "throne room", "трон", "тронный зал"],
        "cues": ["father", "sebastian", "power", "корон", "себастьян", "зал"],
        "min_cues": 1,
    },
    "sharuan_forest_bridge": {
        "title": "Sharuan Forest Bridge",
        "aliases": ["sharuan forest bridge", "forest bridge", "мост", "шаруанский лес"],
        "cues": ["stream", "dream", "лес", "ручей", "сон"],
        "min_cues": 1,
    },
    "paradise_vision": {
        "title": "Paradise Vision",
        "aliases": ["paradise vision", "paradise", "рай", "видение"],
        "cues": ["light", "noctus", "heaven", "свет", "изгнан", "ноктус"],
        "min_cues": 1,
    },
    "krimora_forest": {
        "title": "Krimora Forest",
        "aliases": ["krimora forest", "krimora", "кремор", "кримор"],
        "cues": [
            "crimson",
            "red sky",
            "burgundy",
            "predatory silence",
            "багров",
            "алое небо",
            "красн",
            "лес",
            "клетк",
        ],
        "min_cues": 2,
    },
    "dwarven_caves": {
        "title": "Dwarven Caves",
        "aliases": ["dwarven caves", "dwarven", "гном", "пещер", "дварф"],
        "cues": [
            "stone gate",
            "vault",
            "forge",
            "torch",
            "throne hall",
            "камен",
            "врата",
            "кузн",
            "факел",
            "свод",
            "трон",
        ],
        "min_cues": 2,
    },
    "old_forest": {
        "title": "Old Forest",
        "aliases": ["old forest", "старый лес"],
        "cues": ["battle", "goblin", "очищ", "битв", "гоблин"],
        "min_cues": 1,
    },
    "training_grounds": {
        "title": "Training Grounds",
        "aliases": ["training grounds", "training yard", "трениров"],
        "cues": ["trial", "protocol", "armor", "испытан", "протокол", "брон"],
        "min_cues": 1,
    },
}


def map_runtime_location_to_canonical(raw_token: Any) -> Optional[str]:
    token = _normalize_token(raw_token)
    if not token:
        return None
    if token in _RUNTIME_TO_CANONICAL:
        return _RUNTIME_TO_CANONICAL[token]
    if token.startswith("dwarven_caves"):
        return "dwarven_caves"
    if token.startswith("krimora_forest"):
        return "krimora_forest"
    if token.startswith("kremor_forest"):
        return "krimora_forest"
    return None


def infer_canonical_locations_for_run(plan_name: Any, launcher_location: Any = "") -> List[str]:
    locations: Set[str] = set()
    launcher_canonical = map_runtime_location_to_canonical(launcher_location)
    if launcher_canonical:
        locations.add(launcher_canonical)

    resolved_plan = resolve_video_bot_plan_name(plan_name)
    try:
        events = build_video_bot_events(resolved_plan)
    except Exception:
        events = []
    for row in events:
        if not isinstance(row, dict):
            continue
        kind = _normalize_token(row.get("type", ""))
        if kind not in {"teleport", "portal_jump"}:
            continue
        target = row.get("target", "")
        canonical = map_runtime_location_to_canonical(target)
        if canonical:
            locations.add(canonical)
    return sorted(locations)


def read_docx_text(docx_path: Path) -> str:
    from docx import Document  # Imported lazily for environments without python-docx.

    doc = Document(str(docx_path))
    chunks: List[str] = []
    for paragraph in doc.paragraphs:
        text = str(paragraph.text or "").strip()
        if text:
            chunks.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = str(cell.text or "").strip()
                if text:
                    chunks.append(text)
    return "\n".join(chunks).strip()


def _load_json(path: Path) -> Dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8-sig"))


def load_world_snapshot(project_root: Path) -> Dict[str, Any]:
    layout_path = project_root / "data" / "world" / "layout.json"
    interiors_path = project_root / "data" / "world" / "interiors.json"
    world_source_path = project_root / "src" / "world" / "sharuan_world.py"

    layout = _load_json(layout_path) if layout_path.exists() else {}
    interiors = _load_json(interiors_path) if interiors_path.exists() else {}
    world_source_text = world_source_path.read_text(encoding="utf-8", errors="ignore") if world_source_path.exists() else ""

    zones = set()
    for row in layout.get("zones", []) if isinstance(layout.get("zones"), list) else []:
        if isinstance(row, dict):
            zones.add(_normalize_token(row.get("id", "")))

    door_targets = set()
    for row in layout.get("location_doors", []) if isinstance(layout.get("location_doors"), list) else []:
        if isinstance(row, dict):
            door_targets.add(_normalize_token(row.get("to", "")))

    story_landmarks = set()
    if isinstance(layout.get("story_landmarks"), dict):
        story_landmarks = {_normalize_token(k) for k in layout["story_landmarks"].keys()}

    route_keys = set()
    if isinstance(layout.get("routes"), dict):
        route_keys = {_normalize_token(k) for k in layout["routes"].keys()}

    inner_buildings = set()
    castle_cfg = layout.get("castle", {}) if isinstance(layout.get("castle"), dict) else {}
    for row in castle_cfg.get("inner_buildings", []) if isinstance(castle_cfg.get("inner_buildings"), list) else []:
        if isinstance(row, dict):
            inner_buildings.add(_normalize_token(row.get("type", "")))

    interior_complexes = {_normalize_token(k) for k in interiors.keys()} if isinstance(interiors, dict) else set()
    interior_rooms = set()
    if isinstance(interiors, dict):
        for payload in interiors.values():
            if not isinstance(payload, dict):
                continue
            for room in payload.get("rooms", []) if isinstance(payload.get("rooms"), list) else []:
                if isinstance(room, dict):
                    interior_rooms.add(_normalize_token(room.get("id", "")))

    port_cfg = layout.get("port", {}) if isinstance(layout.get("port"), dict) else {}
    port_counts = {
        "dock_segments": len(port_cfg.get("dock_segments", []) if isinstance(port_cfg.get("dock_segments"), list) else []),
        "market_stalls": len(port_cfg.get("market_stalls", []) if isinstance(port_cfg.get("market_stalls"), list) else []),
        "moored_boats": len(port_cfg.get("moored_boats", []) if isinstance(port_cfg.get("moored_boats"), list) else []),
        "harbor_lanterns": len(port_cfg.get("harbor_lanterns", []) if isinstance(port_cfg.get("harbor_lanterns"), list) else []),
    }

    return {
        "layout_path": str(layout_path),
        "interiors_path": str(interiors_path),
        "world_source_path": str(world_source_path),
        "zones": zones,
        "door_targets": door_targets,
        "story_landmarks": story_landmarks,
        "route_keys": route_keys,
        "inner_buildings": inner_buildings,
        "interior_complexes": interior_complexes,
        "interior_rooms": interior_rooms,
        "port_counts": port_counts,
        "world_source_text": world_source_text,
    }


def _check(id_token: str, ok: bool, details: str) -> Dict[str, Any]:
    return {"id": id_token, "ok": bool(ok), "details": details}


def evaluate_world_checks(canonical_location: str, snapshot: Dict[str, Any]) -> List[Dict[str, Any]]:
    zones = set(snapshot.get("zones", set()))
    door_targets = set(snapshot.get("door_targets", set()))
    story_landmarks = set(snapshot.get("story_landmarks", set()))
    route_keys = set(snapshot.get("route_keys", set()))
    inner_buildings = set(snapshot.get("inner_buildings", set()))
    interior_complexes = set(snapshot.get("interior_complexes", set()))
    interior_rooms = set(snapshot.get("interior_rooms", set()))
    port_counts = dict(snapshot.get("port_counts", {}))
    source_text = _normalize_text(snapshot.get("world_source_text", ""))

    if canonical_location == "port_market":
        return [
            _check("port_dock_segments", int(port_counts.get("dock_segments", 0)) > 0, f"dock_segments={port_counts.get('dock_segments', 0)}"),
            _check("port_market_stalls", int(port_counts.get("market_stalls", 0)) > 0, f"market_stalls={port_counts.get('market_stalls', 0)}"),
            _check("port_moored_boats", int(port_counts.get("moored_boats", 0)) > 0, f"moored_boats={port_counts.get('moored_boats', 0)}"),
            _check("port_harbor_lanterns", int(port_counts.get("harbor_lanterns", 0)) > 0, f"harbor_lanterns={port_counts.get('harbor_lanterns', 0)}"),
        ]

    if canonical_location == "castle_interior":
        return [
            _check("zone_castle_interior", "castle_interior" in zones, "zone 'castle_interior' exists"),
            _check("castle_has_stable", "stable" in inner_buildings, "castle.inner_buildings includes stable"),
            _check("door_target_castle_interior", "castle_interior" in door_targets, "location door target includes castle interior"),
        ]

    if canonical_location == "prince_chamber":
        return [
            _check("zone_prince_chamber", "prince_chamber" in zones, "zone 'prince_chamber' exists"),
            _check("interior_prince_room", "prince_chambers" in interior_rooms, "interior room 'prince_chambers' exists"),
        ]

    if canonical_location == "world_map_gallery":
        return [
            _check("zone_world_map_gallery", "world_map_gallery" in zones, "zone 'world_map_gallery' exists"),
            _check("door_target_world_map_gallery", "world_map_gallery" in door_targets, "location door target includes world map gallery"),
        ]

    if canonical_location == "royal_laundry":
        return [
            _check("zone_royal_laundry", "royal_laundry" in zones, "zone 'royal_laundry' exists"),
        ]

    if canonical_location == "throne_hall":
        return [
            _check("zone_throne_hall", "throne_hall" in zones, "zone 'throne_hall' exists"),
            _check("interior_great_hall", "great_hall" in interior_rooms, "interior room 'great_hall' exists"),
            _check("castle_torch_lighting", ("castle_torch" in source_text) or ("torch" in source_text and "throne hall" in source_text), "world source includes throne torch lighting"),
        ]

    if canonical_location == "sharuan_forest_bridge":
        return [
            _check("zone_sharuan_forest_bridge", "sharuan_forest_bridge" in zones, "zone 'sharuan_forest_bridge' exists"),
            _check("story_landmark_sharuan_bridge", "sharuan_bridge" in story_landmarks, "story landmark 'sharuan_bridge' exists"),
        ]

    if canonical_location == "paradise_vision":
        return [
            _check("zone_paradise_vision", "paradise_vision" in zones, "zone 'paradise_vision' exists"),
        ]

    if canonical_location == "krimora_forest":
        return [
            _check("zone_krimora_crash", "krimora_forest_crash" in zones, "zone 'krimora_forest_crash' exists"),
            _check("zone_krimora_cage", "krimora_forest_cage" in zones, "zone 'krimora_forest_cage' exists"),
            _check("story_landmark_adrian_cage", "adrian_cage" in story_landmarks, "story landmark 'adrian_cage' exists"),
            _check("krimora_story_setpiece", "krimora cage clearing" in source_text, "world source includes Krimora cage clearing"),
        ]

    if canonical_location == "dwarven_caves":
        return [
            _check("zone_dwarven_gate", "dwarven_caves_gate" in zones, "zone 'dwarven_caves_gate' exists"),
            _check("zone_dwarven_halls", "dwarven_caves_halls" in zones, "zone 'dwarven_caves_halls' exists"),
            _check("zone_dwarven_throne", "dwarven_caves_throne" in zones, "zone 'dwarven_caves_throne' exists"),
            _check("interior_dwarven_caves", "dwarven_caves" in interior_complexes, "interior complex 'dwarven_caves' exists"),
            _check("dwarven_gate_landmark", "dwarven_gate" in story_landmarks, "story landmark 'dwarven_gate' exists"),
            _check("dwarven_no_sky_shell", ("block sky" in source_text) and ("dwarven" in source_text), "world source blocks cave sky"),
            _check("dwarven_torch_route", ("spawn_torch" in source_text) and ("dwarven" in source_text), "world source has dwarven torch route"),
        ]

    if canonical_location == "old_forest":
        return [
            _check("zone_old_forest", "old_forest" in zones, "zone 'old_forest' exists"),
            _check("route_forest_zone", "forest_zone" in route_keys, "layout.routes has forest_zone"),
        ]

    if canonical_location == "training_grounds":
        return [
            _check("zone_training_grounds", "training_grounds" in zones, "zone 'training_grounds' exists"),
        ]

    return []


def _count_hits(text: str, term: str) -> int:
    token = _normalize_text(term)
    if not token:
        return 0
    return int(text.count(token))


def _build_excerpt(text: str, terms: Sequence[str], radius: int = 180) -> str:
    first_index = -1
    for term in terms:
        token = _normalize_text(term)
        if not token:
            continue
        idx = text.find(token)
        if idx >= 0 and (first_index < 0 or idx < first_index):
            first_index = idx
    if first_index < 0:
        return ""
    start = max(0, first_index - radius)
    end = min(len(text), first_index + radius)
    snippet = text[start:end].strip()
    if start > 0:
        snippet = "..." + snippet
    if end < len(text):
        snippet = snippet + "..."
    return snippet


def evaluate_book_support(book_text: str, canonical_location: str) -> Dict[str, Any]:
    profile = BOOK_LOCATION_PROFILES.get(canonical_location)
    if not profile:
        return {
            "ok": False,
            "reason": "no_profile",
            "title": canonical_location,
            "alias_hits": [],
            "cue_hits": [],
            "min_cues": 0,
            "excerpt": "",
        }
    normalized = _normalize_text(book_text)
    alias_terms = list(profile.get("aliases", []))
    cue_terms = list(profile.get("cues", []))

    alias_hits = [{"term": term, "hits": _count_hits(normalized, term)} for term in alias_terms]
    cue_hits = [{"term": term, "hits": _count_hits(normalized, term)} for term in cue_terms]
    alias_total = sum(int(row["hits"]) for row in alias_hits)
    cue_present = sum(1 for row in cue_hits if int(row["hits"]) > 0)
    min_cues = int(profile.get("min_cues", 1) or 1)

    ok = alias_total > 0 and cue_present >= min_cues
    excerpt = _build_excerpt(normalized, alias_terms + cue_terms)
    return {
        "ok": bool(ok),
        "reason": "ok" if ok else "book_evidence_missing",
        "title": str(profile.get("title", canonical_location)),
        "alias_hits": alias_hits,
        "cue_hits": cue_hits,
        "alias_total": alias_total,
        "cue_present": cue_present,
        "min_cues": min_cues,
        "excerpt": excerpt,
    }


def evaluate_location_conformance(
    canonical_location: str,
    book_text: str,
    snapshot: Dict[str, Any],
) -> Dict[str, Any]:
    book = evaluate_book_support(book_text, canonical_location)
    world_checks = evaluate_world_checks(canonical_location, snapshot)
    world_ok = all(bool(row.get("ok", False)) for row in world_checks) if world_checks else False

    if not book.get("ok", False):
        status = "fail"
        reason = "book_not_confirmed"
    elif not world_ok:
        status = "fail"
        reason = "world_checks_failed"
    else:
        status = "pass"
        reason = "ok"

    return {
        "canonical_location": canonical_location,
        "title": BOOK_LOCATION_PROFILES.get(canonical_location, {}).get("title", canonical_location),
        "status": status,
        "reason": reason,
        "book": book,
        "world_checks": world_checks,
        "world_ok": world_ok,
    }
